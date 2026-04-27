"""
File parsing service for extracting text from PDF and DOCX files.
"""

import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using pdfplumber.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the PDF cannot be parsed or is empty.
    """
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise ValueError("PDF file has no pages.")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts).strip()
        if not full_text:
            raise ValueError("No readable text found in the PDF. It may be image-based.")
        return full_text

    except ValueError:
        raise
    except Exception as exc:
        logger.error("Failed to parse PDF: %s", exc)
        raise ValueError(f"Could not read PDF file: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the DOCX cannot be parsed or is empty.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs).strip()
        if not full_text:
            raise ValueError("No readable text found in the DOCX file.")
        return full_text

    except ValueError:
        raise
    except Exception as exc:
        logger.error("Failed to parse DOCX: %s", exc)
        raise ValueError(f"Could not read DOCX file: {exc}") from exc


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route file to the correct parser based on extension.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename (used to determine file type).

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file type is unsupported or parsing fails.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif suffix == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{suffix}'. Please upload a PDF or DOCX.")
